from collections import defaultdict
from html import escape
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image as ReportLabImage, Paragraph, SimpleDocTemplate, Spacer

from releases.reporting import support_line


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 117, 109)
DOCX_LATIN_FONT = "Calibri"
DOCX_KOREAN_FONT = "NanumGothic"
NANUM_FONT_PATHS = (
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = DOCX_LATIN_FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_KOREAN_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_hyperlink(paragraph, text, url):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), DOCX_LATIN_FONT)
    fonts.set(qn("w:hAnsi"), DOCX_LATIN_FONT)
    fonts.set(qn("w:eastAsia"), DOCX_KOREAN_FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "176B4D")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend((fonts, color, underline))
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.extend((run_properties, text_element))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_docx_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = DOCX_LATIN_FONT
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_KOREAN_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[name]
        style.font.name = DOCX_LATIN_FONT
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), DOCX_KOREAN_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def render_docx_bytes(report):
    document = Document()
    configure_docx_styles(document)
    section = document.sections[0]
    branding = report["branding"]

    header = section.header.paragraphs[0]
    header.text = "POSTGRESQL UPGRADE BRIEF · APPROVED REPORT"
    set_run_font(header.runs[0], size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Page ")
    add_page_field(footer)
    for run in footer.runs:
        set_run_font(run, size=9, color=MUTED)

    if branding["logo_path"]:
        logo = document.add_paragraph()
        logo.paragraph_format.space_after = Pt(10)
        logo.add_run().add_picture(branding["logo_path"], width=Inches(1.35))
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    set_run_font(kicker.add_run("POSTGRESQL OFFICIAL RELEASE INTELLIGENCE"), size=9, bold=True, color=BLUE)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_run_font(title.add_run(report["title"]), size=24, bold=True, color=RGBColor(24, 37, 31))
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("고객 전달용 승인 보고서" if report["level"] == "customer" else "DBA 기술 검토용 승인 보고서"), size=12, color=MUTED)

    summary = report["summary"]
    document.add_heading("보고서 개요", level=1)
    summary_fields = [
        ("업그레이드 범위", f"{summary['from_version']} (제외) → {summary['to_version']} (포함)"),
        ("포함 릴리스", f"{summary['release_count']}개"),
        ("승인 항목", f"{len(report['items'])}개"),
        ("생성시각", report["generated_at"]),
    ]
    if branding["customer_name"]:
        summary_fields.append(("고객명", branding["customer_name"]))
    if branding["project_name"]:
        summary_fields.append(("프로젝트명", branding["project_name"]))
    summary_fields.extend([
        ("AS-IS 지원", support_line("AS-IS", summary["from_support"]).removeprefix("AS-IS: ")),
        ("TO-BE 지원", support_line("TO-BE", summary["to_support"]).removeprefix("TO-BE: ")),
    ])
    for label, value in summary_fields:
        paragraph = document.add_paragraph()
        set_run_font(paragraph.add_run(f"{label}: "), bold=True)
        set_run_font(paragraph.add_run(value))

    document.add_heading("승인된 변경사항", level=1)
    grouped = defaultdict(list)
    for item in report["items"]:
        grouped[item.version].append(item)
    for version, items in grouped.items():
        document.add_heading(f"PostgreSQL {version}", level=2)
        for item in items:
            heading = document.add_paragraph()
            heading.paragraph_format.keep_with_next = True
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
            set_run_font(heading.add_run(f"{item.area} · {item.change_type}"), size=11, bold=True, color=DARK_BLUE)
            body = document.add_paragraph(item.text)
            body.paragraph_format.keep_together = True
            source = document.add_paragraph()
            source.paragraph_format.space_after = Pt(8)
            add_hyperlink(source, "PostgreSQL 공식 원문", item.source_url)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def register_pdf_font():
    for font_path in NANUM_FONT_PATHS:
        if Path(font_path).exists():
            pdfmetrics.registerFont(TTFont("NanumGothic", font_path))
            return "NanumGothic"
    raise RuntimeError("NanumGothic font not found")


def render_pdf_bytes(report):
    font_name = register_pdf_font()
    output = BytesIO()

    def draw_page(canvas, document):
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(HexColor("#66756D"))
        canvas.drawString(inch, 10.45 * inch, "POSTGRESQL UPGRADE BRIEF · APPROVED REPORT")
        canvas.drawRightString(7.5 * inch, 0.55 * inch, f"Page {document.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        output,
        pagesize=LETTER,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=0.9 * inch,
        bottomMargin=0.8 * inch,
        title=report["title"],
        author="PostgreSQL Upgrade Brief Generator",
    )
    styles = getSampleStyleSheet()
    branding = report["branding"]
    body = ParagraphStyle("BodyKR", parent=styles["BodyText"], fontName=font_name, fontSize=10, leading=13.2, spaceAfter=6, textColor=HexColor("#18251F"))
    title = ParagraphStyle("TitleKR", parent=body, fontSize=23, leading=28, spaceAfter=5, textColor=HexColor("#18251F"))
    subtitle = ParagraphStyle("SubtitleKR", parent=body, fontSize=11, leading=14, spaceAfter=18, textColor=HexColor("#66756D"))
    h1 = ParagraphStyle("H1KR", parent=body, fontSize=15, leading=19, spaceBefore=16, spaceAfter=8, textColor=HexColor("#2E74B5"), keepWithNext=True)
    h2 = ParagraphStyle("H2KR", parent=body, fontSize=12, leading=15, spaceBefore=12, spaceAfter=6, textColor=HexColor("#2E74B5"), keepWithNext=True)
    label = ParagraphStyle("LabelKR", parent=body, fontSize=10, leading=13, spaceBefore=7, spaceAfter=3, textColor=HexColor("#1F4D78"), keepWithNext=True)
    source_style = ParagraphStyle("SourceKR", parent=body, fontSize=9, leading=12, spaceAfter=8, textColor=HexColor("#176B4D"))

    summary = report["summary"]
    story = []
    if branding["logo_path"]:
        logo_width, logo_height = ImageReader(branding["logo_path"]).getSize()
        logo_scale = min(1.35 * inch / logo_width, 0.58 * inch / logo_height)
        report_logo = ReportLabImage(branding["logo_path"], width=logo_width * logo_scale, height=logo_height * logo_scale)
        report_logo.hAlign = "LEFT"
        story.extend([report_logo, Spacer(1, 8)])
    story.extend([
        Paragraph("POSTGRESQL OFFICIAL RELEASE INTELLIGENCE", label),
        Paragraph(escape(report["title"]), title),
        Paragraph("고객 전달용 승인 보고서" if report["level"] == "customer" else "DBA 기술 검토용 승인 보고서", subtitle),
        Paragraph("보고서 개요", h1),
    ])
    summary_fields = [
        ("업그레이드 범위", f"{summary['from_version']} (제외) → {summary['to_version']} (포함)"),
        ("포함 릴리스", f"{summary['release_count']}개"),
        ("승인 항목", f"{len(report['items'])}개"),
        ("생성시각", report["generated_at"]),
    ]
    if branding["customer_name"]:
        summary_fields.append(("고객명", branding["customer_name"]))
    if branding["project_name"]:
        summary_fields.append(("프로젝트명", branding["project_name"]))
    summary_fields.extend([
        ("AS-IS 지원", support_line("AS-IS", summary["from_support"]).removeprefix("AS-IS: ")),
        ("TO-BE 지원", support_line("TO-BE", summary["to_support"]).removeprefix("TO-BE: ")),
    ])
    for key, value in summary_fields:
        story.append(Paragraph(f"<b>{escape(key)}:</b> {escape(value)}", body))
    story.append(Paragraph("승인된 변경사항", h1))

    grouped = defaultdict(list)
    for item in report["items"]:
        grouped[item.version].append(item)
    for version, items in grouped.items():
        story.append(Paragraph(f"PostgreSQL {escape(version)}", h2))
        for item in items:
            story.append(Paragraph(f"{escape(item.area)} · {escape(item.change_type)}", label))
            story.append(Paragraph(escape(item.text).replace("\n", "<br/>"), body))
            story.append(Paragraph(f'<link href="{escape(item.source_url, quote=True)}">PostgreSQL 공식 원문</link>', source_style))
    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return output.getvalue()
